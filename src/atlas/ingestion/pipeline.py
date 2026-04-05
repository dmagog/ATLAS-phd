"""
Ingestion pipeline: accept → extract → normalize → chunk → embed → index

Supported formats:
  - PDF, DOCX, TXT, MD  — plain text extraction, page number unknown
  - JSONL               — page-aware format: {"page": N, "text": "...", "source_pdf": "..."}
                          Preserves page numbers in chunk metadata for accurate citations.
"""
import asyncio
import hashlib
import io
import json
import re
import uuid
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

import httpx
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from atlas.core.config import settings
from atlas.core.logging import logger
from atlas.db.models import Chunk, Document, IngestionJob


SUPPORTED_MIME_TYPES = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
    "text/markdown": ".md",
    "application/jsonl": ".jsonl",
}

# Also accept by extension regardless of MIME (browsers often send text/plain for .jsonl)
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".jsonl"}

CHUNK_MIN = 800
CHUNK_MAX = 1200
CHUNK_OVERLAP = 150


@dataclass
class ChunkMeta:
    text: str
    page: Optional[int] = None
    section: Optional[str] = None


@dataclass
class FileResult:
    filename: str
    status: str  # "accepted" | "rejected" | "processed" | "failed"
    reason: Optional[str] = None
    document_id: Optional[str] = None
    chunks_created: int = 0


@dataclass
class RawFile:
    filename: str
    content: bytes
    mime_type: str


# ── Stage 1: Accept ──────────────────────────────────────────────────────────

def _is_supported(raw: RawFile) -> bool:
    ext = Path(raw.filename).suffix.lower()
    return raw.mime_type in SUPPORTED_MIME_TYPES or ext in SUPPORTED_EXTENSIONS


def accept_file(raw: RawFile) -> tuple[bool, str]:
    """Validate format by MIME type or file extension."""
    if not _is_supported(raw):
        return False, "UNSUPPORTED_FORMAT"
    if len(raw.content) == 0:
        return False, "EMPTY_FILE"
    return True, ""


def compute_sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _is_jsonl(raw: RawFile) -> bool:
    return Path(raw.filename).suffix.lower() == ".jsonl"


# ── Stage 2: Extract ─────────────────────────────────────────────────────────

def extract_pages(raw: RawFile) -> list[ChunkMeta]:
    """
    Extract content as a list of ChunkMeta items.
    For JSONL: one item per page (with page number).
    For other formats: one item with the full text (page=None).
    """
    if _is_jsonl(raw):
        return _extract_jsonl(raw.content)
    text = _extract_plain(raw)
    return [ChunkMeta(text=text)]


def _extract_jsonl(content: bytes) -> list[ChunkMeta]:
    pages = []
    for line in content.decode("utf-8", errors="replace").splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            obj = json.loads(line)
            text = obj.get("text", "").strip()
            page = obj.get("page")
            if text:
                pages.append(ChunkMeta(text=text, page=page))
        except json.JSONDecodeError:
            continue
    return pages


def _extract_plain(raw: RawFile) -> str:
    if raw.mime_type == "application/pdf":
        return _extract_pdf(raw.content)
    if raw.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(raw.content)
    return raw.content.decode("utf-8", errors="replace")


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(content))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── Stage 3: Normalize ───────────────────────────────────────────────────────

_SUSPICIOUS = re.compile(r"(ignore previous instructions|system prompt|disregard)", re.IGNORECASE)


def normalize(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def has_suspicious_patterns(text: str) -> bool:
    return bool(_SUSPICIOUS.search(text))


# ── Stage 4: Chunk ───────────────────────────────────────────────────────────

def chunk_pages(pages: list[ChunkMeta]) -> list[ChunkMeta]:
    """
    Chunk a list of page-level items into retrieval-sized chunks.

    Strategy:
      - Accumulate pages until CHUNK_MAX is reached, then emit a chunk.
      - The chunk inherits the page number of the FIRST page it started from.
      - If a single page exceeds CHUNK_MAX, split it by sentences.
      - Add CHUNK_OVERLAP of text from the previous chunk as prefix.
    """
    result: list[ChunkMeta] = []
    current_text = ""
    current_page: Optional[int] = None
    prev_tail = ""

    def _emit(text: str, page: Optional[int]) -> None:
        nonlocal prev_tail
        if text:
            result.append(ChunkMeta(text=text.strip(), page=page))
            prev_tail = text[-CHUNK_OVERLAP:]

    for item in pages:
        text = normalize(item.text)
        if not text:
            continue

        if has_suspicious_patterns(text):
            logger.warning("suspicious_content_in_page", page=item.page)

        # If adding this page exceeds limit — emit current and start fresh
        if current_text and len(current_text) + len(text) + 2 > CHUNK_MAX:
            _emit(current_text, current_page)
            overlap = prev_tail if prev_tail else ""
            current_text = (overlap + "\n\n" + text).strip() if overlap else text
            current_page = item.page
        else:
            # Merge into current chunk
            if not current_text:
                overlap = prev_tail if prev_tail else ""
                current_text = (overlap + "\n\n" + text).strip() if overlap else text
                current_page = item.page
            else:
                current_text = (current_text + "\n\n" + text).strip()

        # Oversized single page — flush immediately in sentence-sized pieces
        while len(current_text) > CHUNK_MAX:
            pieces = _split_long_text(current_text)
            for piece in pieces[:-1]:
                _emit(piece, current_page)
            current_text = pieces[-1] if pieces else ""

    if current_text:
        _emit(current_text, current_page)

    return result


def _split_long_text(text: str) -> list[str]:
    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks: list[str] = []
    current = ""
    for sent in sentences:
        # Sentence itself exceeds limit (formulas, tables, no punctuation) — force-split by chars
        if len(sent) > CHUNK_MAX:
            if current:
                chunks.append(current)
                current = ""
            for i in range(0, len(sent), CHUNK_MAX):
                chunks.append(sent[i: i + CHUNK_MAX])
        elif len(current) + len(sent) + 1 <= CHUNK_MAX:
            current = (current + " " + sent).strip()
        else:
            if current:
                chunks.append(current)
            current = sent
    if current:
        chunks.append(current)
    return chunks if chunks else [text[:CHUNK_MAX]]


# ── Thread-safe extract+chunk helper ─────────────────────────────────────────

def _extract_and_chunk(raw: RawFile) -> tuple[list[ChunkMeta], list[ChunkMeta]]:
    """Synchronous extract + chunk — safe to run in asyncio.to_thread."""
    pages = extract_pages(raw)
    if not pages:
        return [], []
    chunks = chunk_pages(pages)
    return pages, chunks


# ── Stage 5: Embed ───────────────────────────────────────────────────────────

async def embed_chunks(texts: list[str]) -> list[list[float]]:
    async with httpx.AsyncClient(timeout=60.0) as client:
        resp = await client.post(
            f"{settings.embeddings_url}/embed",
            json={"texts": texts},
        )
        resp.raise_for_status()
        return resp.json()["embeddings"]


# ── Stage 6: Index ───────────────────────────────────────────────────────────

async def index_document(
    db: AsyncSession,
    raw: RawFile,
    sha256: str,
    chunks: list[ChunkMeta],
    embeddings: list[list[float]],
    file_path: str,
    job_id: str,
) -> Document:
    doc = Document(
        id=uuid.uuid4(),
        title=Path(raw.filename).stem,
        filename=raw.filename,
        sha256=sha256,
        file_path=file_path,
        mime_type=raw.mime_type,
    )
    db.add(doc)
    await db.flush()

    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        db.add(Chunk(
            id=uuid.uuid4(),
            document_id=doc.id,
            chunk_index=i,
            text=chunk.text,
            page=chunk.page,
            section=chunk.section,
            embedding=embedding,
        ))

    await db.commit()
    logger.info("document_indexed", document_id=str(doc.id), filename=raw.filename,
                chunks=len(chunks), job_id=job_id)
    return doc


# ── Orchestrator ─────────────────────────────────────────────────────────────

async def process_file(
    db: AsyncSession,
    raw: RawFile,
    corpus_dir: Path,
    job_id: str,
) -> FileResult:
    # Stage 1: Accept
    ok, reason = accept_file(raw)
    if not ok:
        return FileResult(filename=raw.filename, status="rejected", reason=reason)

    # Idempotency by SHA-256
    sha256 = compute_sha256(raw.content)
    existing = await db.execute(select(Document).where(Document.sha256 == sha256))
    if existing.scalar_one_or_none():
        return FileResult(filename=raw.filename, status="rejected", reason="DUPLICATE")

    try:
        # Stage 2+3+4: Extract / Normalize / Chunk — run in thread to not block event loop
        pages, chunks = await asyncio.to_thread(_extract_and_chunk, raw)
        if not pages:
            return FileResult(filename=raw.filename, status="rejected", reason="EMPTY_CONTENT")
        if not chunks:
            return FileResult(filename=raw.filename, status="failed", reason="CHUNKING_FAILED")

        # Stage 5: Embed in batches of 64, yielding between batches
        all_embeddings: list[list[float]] = []
        batch_size = 64
        texts = [c.text for c in chunks]
        for i in range(0, len(texts), batch_size):
            all_embeddings.extend(await embed_chunks(texts[i: i + batch_size]))
            await asyncio.sleep(0)  # yield to event loop between batches

        # Save raw file to corpus (in thread — file I/O)
        file_path = str(corpus_dir / raw.filename)
        await asyncio.to_thread((corpus_dir / raw.filename).write_bytes, raw.content)

        # Stage 6: Index
        doc = await index_document(db, raw, sha256, chunks, all_embeddings, file_path, job_id)

        return FileResult(
            filename=raw.filename,
            status="processed",
            document_id=str(doc.id),
            chunks_created=len(chunks),
        )

    except Exception as exc:
        logger.error("ingestion_file_failed", filename=raw.filename, error=str(exc), job_id=job_id)
        return FileResult(filename=raw.filename, status="failed", reason="PROCESSING_ERROR")


async def run_ingestion_job(
    db: AsyncSession,
    job: IngestionJob,
    files: list[RawFile],
    corpus_dir: Path,
) -> IngestionJob:
    job.status = "running"
    job.accepted_files = []
    job.rejected_files = []
    await db.commit()

    accepted: list[dict] = []
    rejected: list[dict] = []
    total = len(files)

    for idx, raw in enumerate(files, 1):
        logger.info("ingestion_file_start", filename=raw.filename,
                    progress=f"{idx}/{total}", job_id=str(job.id))

        result = await process_file(db, raw, corpus_dir, str(job.id))
        entry = {"filename": result.filename, "reason": result.reason}

        if result.status in ("processed", "accepted"):
            entry["document_id"] = result.document_id
            entry["chunks_created"] = result.chunks_created
            accepted.append(entry)
        else:
            rejected.append(entry)

        # Commit partial progress after every file so polling sees live state
        from sqlalchemy.orm.attributes import flag_modified
        job.accepted_files = list(accepted)
        job.rejected_files = list(rejected)
        flag_modified(job, "accepted_files")
        flag_modified(job, "rejected_files")
        await db.commit()

        logger.info("ingestion_file_done", filename=raw.filename,
                    status=result.status, progress=f"{idx}/{total}", job_id=str(job.id))

    job.status = "completed" if accepted and not rejected else (
        "completed_with_errors" if accepted else "failed"
    )
    job.completed_at = datetime.utcnow()
    await db.commit()

    logger.info("ingestion_job_done", job_id=str(job.id), status=job.status,
                accepted=len(accepted), rejected=len(rejected))
    return job
